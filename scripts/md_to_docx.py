#!/usr/bin/env python3
"""将 Markdown 文件转换为格式化的 Word 文档 (.docx)"""

import sys
from pathlib import Path

try:
    import mistune
except ImportError:
    print("请先安装 mistune: pip3 install mistune")
    sys.exit(1)

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import parse_xml
except ImportError:
    print("请先安装 python-docx: pip3 install python-docx")
    sys.exit(1)


def extract_text(children):
    """从 mistune AST children 中提取纯文本"""
    result = []
    for child in children:
        t = child.get('type', '')
        if t == 'text':
            result.append(child.get('raw', ''))
        elif t == 'codespan':
            result.append('`' + child.get('raw', '') + '`')
        elif t == 'strong':
            result.append('**' + extract_text(child.get('children', [])) + '**')
        elif t == 'emphasis':
            result.append('*' + extract_text(child.get('children', [])) + '*')
        elif t == 'link':
            result.append(extract_text(child.get('children', [])))
        elif t == 'linebreak':
            result.append('\n')
        elif 'children' in child:
            result.append(extract_text(child['children']))
    return ''.join(result)


def add_rich_text(para, children):
    """将 mistune inline AST 渲染为带格式的 docx runs"""
    for child in children:
        t = child.get('type', '')
        if t == 'text':
            para.add_run(child.get('raw', ''))
        elif t == 'codespan':
            run = para.add_run(child.get('raw', ''))
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xC4, 0x1E, 0x3A)
        elif t == 'strong':
            run = para.add_run(extract_text(child.get('children', [])))
            run.bold = True
        elif t == 'emphasis':
            run = para.add_run(extract_text(child.get('children', [])))
            run.italic = True
        elif t == 'link':
            run = para.add_run(extract_text(child.get('children', [])))
            run.font.color.rgb = RGBColor(0x33, 0x70, 0xFF)
            run.underline = True
        elif t == 'linebreak':
            para.add_run('\n')
        elif 'children' in child:
            add_rich_text(para, child['children'])


def render_ast_to_docx(doc, ast_nodes):
    """将 mistune AST 节点渲染到 docx 文档"""
    for node in ast_nodes:
        node_type = node.get('type', '')

        if node_type == 'heading':
            level = node.get('attrs', {}).get('level', 1)
            text = extract_text(node.get('children', []))
            heading = doc.add_heading(level=min(level, 6))
            heading.clear()
            run = heading.add_run(text)
            run.font.name = 'Microsoft YaHei'
            run.font.size = Pt({1: 18, 2: 16, 3: 14, 4: 12, 5: 11, 6: 10}.get(level, 11))
            if level <= 2:
                run.font.color.rgb = RGBColor(0x33, 0x70, 0xFF)
                run.bold = True
            heading.paragraph_format.space_before = Pt(12)
            heading.paragraph_format.space_after = Pt(6)

        elif node_type == 'paragraph':
            children = node.get('children', [])
            if not children:
                continue
            para = doc.add_paragraph()
            para.paragraph_format.space_after = Pt(4)
            para.paragraph_format.line_spacing = 1.15
            add_rich_text(para, children)

        elif node_type == 'block_code':
            code = node.get('raw', '')
            para = doc.add_paragraph()
            para.paragraph_format.space_after = Pt(4)
            para.paragraph_format.left_indent = Inches(0.2)
            run = para.add_run(code)
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

        elif node_type == 'block_quote':
            children = node.get('children', [])
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Inches(0.3)
            add_rich_text(para, children)
            for run in para.runs:
                run.italic = True
                run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        elif node_type == 'list':
            ordered = node.get('ordered', False)
            items = node.get('children', [])
            for idx, item in enumerate(items, 1):
                item_children = item.get('children', [])
                para = doc.add_paragraph(style='List Bullet' if not ordered else 'List Number')
                para.paragraph_format.space_after = Pt(2)
                para.paragraph_format.left_indent = Inches(0.3)
                add_rich_text(para, item_children)

        elif node_type == 'table':
            # mistune v3 table AST:
            # children[0] = table_head (contains table_cell list)
            # children[1] = table_body (contains table_row list)
            # each table_row contains table_cell list
            table_children = node.get('children', [])
            headers = []
            rows = []

            for tc in table_children:
                tc_type = tc.get('type', '')
                if tc_type == 'table_head':
                    for cell in tc.get('children', []):
                        headers.append(extract_text(cell.get('children', [])))
                elif tc_type == 'table_body':
                    for row in tc.get('children', []):
                        cells = []
                        for cell in row.get('children', []):
                            cells.append(extract_text(cell.get('children', [])))
                        rows.append(cells)

            total_rows = (1 if headers else 0) + len(rows)
            total_cols = len(headers) if headers else (len(rows[0]) if rows else 1)
            if total_rows == 0 or total_cols == 0:
                continue

            table = doc.add_table(rows=total_rows, cols=total_cols)
            table.style = 'Table Grid'
            row_idx = 0
            if headers:
                for col_idx, header_text in enumerate(headers):
                    if col_idx >= total_cols:
                        break
                    cell = table.rows[row_idx].cells[col_idx]
                    cell.text = header_text
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                            run.font.size = Pt(10)
                            run.font.name = 'Microsoft YaHei'
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    shading = parse_xml(r'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="F1F5F9"/>')
                    cell._tc.get_or_add_tcPr().append(shading)
                row_idx += 1
            for row_data in rows:
                for col_idx, cell_text in enumerate(row_data):
                    if col_idx >= total_cols:
                        break
                    cell = table.rows[row_idx].cells[col_idx]
                    cell.text = cell_text
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(10)
                            run.font.name = 'Microsoft YaHei'
                row_idx += 1
            doc.add_paragraph()

        elif node_type == 'blank_line':
            pass

        elif node_type == 'thematic_break':
            para = doc.add_paragraph()
            para.paragraph_format.space_before = Pt(6)
            para.paragraph_format.space_after = Pt(6)
            run = para.add_run('─' * 40)
            run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
            run.font.size = Pt(8)


def convert(md_path, docx_path):
    md_content = Path(md_path).read_text(encoding='utf-8')
    md = mistune.create_markdown(renderer=None, plugins=['table'])
    ast_tuple = md.parse(md_content)
    ast = ast_tuple[0] if isinstance(ast_tuple, tuple) else ast_tuple

    doc = Document()

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Microsoft YaHei'
    font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    sections = doc.sections[0]
    sections.top_margin = Inches(1)
    sections.bottom_margin = Inches(1)
    sections.left_margin = Inches(1)
    sections.right_margin = Inches(1)

    render_ast_to_docx(doc, ast)

    doc.save(docx_path)
    print(f"已生成: {docx_path}")


if __name__ == '__main__':
    md_file = sys.argv[1] if len(sys.argv) > 1 else 'docs/product-features.md'
    out_file = sys.argv[2] if len(sys.argv) > 2 else 'docs/product-features.docx'
    convert(md_file, out_file)
