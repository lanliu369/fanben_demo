from __future__ import annotations

import argparse
import re
from typing import Optional
import shutil
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

# Word 默认「标题」样式多为主题蓝；统一改为黑色避免印刷偏色
_HEADING_BLACK = RGBColor(0x00, 0x00, 0x00)


_ROOT = Path(__file__).resolve().parent.parent
DEFAULT_SOURCE_MD = _ROOT / "PRD_v1.0.md"
DEFAULT_TARGET_DOCX = _ROOT / "PRD_v1.0_印刷版.docx"
DEFAULT_COVER_VERSION = "v1.0"
DEFAULT_COVER_DATE = "2026年4月"
TARGET_DOCX_FINAL_ALIAS = (
    _ROOT
    / "招标文件范本编制工具平台_产品需求规格说明书_同结构版_v1.0_国能格式_终稿_页面层级版_v7.docx"
)


def set_run_font(run, size=12, bold=False, east_asia="宋体", color_rgb: Optional[RGBColor] = None):
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    if color_rgb is not None:
        run.font.color.rgb = color_rgb


def set_paragraph_line_spacing(paragraph, before=6, after=6, line=1.5):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line


def add_toc(paragraph):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r'TOC \o "1-3" \h \z \u'
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "目录（右键更新域）"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(text)
    run._r.append(fld_end)


def clean_inline_md(text: str) -> str:
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"\*([^*]+)\*", r"\1", text)
    return text.strip()


def parse_headings(lines: list[str]):
    heading_lines = []
    for line in lines:
        m = re.match(r"^(#{1,6})\s+(.*)$", line.strip())
        if m:
            level = len(m.group(1))
            title = clean_inline_md(m.group(2))
            heading_lines.append((level, title))
    return heading_lines


def build_cover(doc: Document, version: str = DEFAULT_COVER_VERSION, cover_date: str = DEFAULT_COVER_DATE):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("招标文件范本编制工具平台")
    set_run_font(r, size=24, bold=True, east_asia="黑体")
    set_paragraph_line_spacing(p, before=120, after=24, line=1.2)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("产品需求规格说明书")
    set_run_font(r2, size=22, bold=True, east_asia="黑体")

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run("（印刷版）")
    set_run_font(r3, size=18, bold=False, east_asia="黑体")
    set_paragraph_line_spacing(p3, before=10, after=60, line=1.2)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rm = meta.add_run(f"版本：{version}")
    set_run_font(rm, size=14, east_asia="宋体")

    meta2 = doc.add_paragraph()
    meta2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rm2 = meta2.add_run(f"编制日期：{cover_date}")
    set_run_font(rm2, size=14, east_asia="宋体")


def apply_header_footer(section):
    section.different_first_page_header_footer = True

    header = section.header
    hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hr = hp.add_run("招标文件范本编制工具平台 产品需求规格说明书（印刷版）")
    set_run_font(hr, size=10, east_asia="宋体")

    footer = section.footer
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run("第 ")
    set_run_font(fr, size=10, east_asia="宋体")

    page_field_run = fp.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    page_field_run._r.append(fld_begin)
    page_field_run._r.append(instr)
    page_field_run._r.append(fld_sep)
    page_field_run._r.append(text)
    page_field_run._r.append(fld_end)

    fr2 = fp.add_run(" 页")
    set_run_font(fr2, size=10, east_asia="宋体")


def add_content(doc: Document, lines: list[str]):
    counters = [0, 0, 0, 0, 0, 0]
    # Word「编号」样式会跨段落延续序号；各小节改用正文自管序号，遇标题重置
    list_serial = 0

    for raw in lines:
        line = raw.rstrip("\n")
        if not line.strip() or line.strip() == "---":
            doc.add_paragraph("")
            continue

        hm = re.match(r"^(#{1,6})\s+(.*)$", line.strip())
        if hm:
            list_serial = 0
            md_level = len(hm.group(1))
            title = clean_inline_md(hm.group(2))

            # PRD 正文不再输出单独「1 全书标题」：封面已有书名，跳过 Markdown 唯一 `#`
            # 原 `##`→Word 一级标题（1、2…），原 `###`→二级（1.1…），整体抬一级
            if md_level == 1:
                continue

            word_level = md_level - 1
            counters[word_level - 1] += 1
            for i in range(word_level, 6):
                counters[i] = 0

            numbering = ".".join(str(x) for x in counters[:word_level] if x > 0)
            text = f"{numbering} {title}"

            p = doc.add_paragraph(text, style=f"Heading {word_level}")
            east = "黑体" if word_level <= 2 else "宋体"
            for run in p.runs:
                set_run_font(
                    run,
                    size=max(16 - word_level, 12),
                    bold=True,
                    east_asia=east,
                    color_rgb=_HEADING_BLACK,
                )
            set_paragraph_line_spacing(p, before=8, after=4, line=1.3)
            continue

        bm = re.match(r"^\s*[-*]\s+(.*)$", line)
        if bm:
            p = doc.add_paragraph(style="List Bullet")
            r = p.add_run(clean_inline_md(bm.group(1)))
            set_run_font(r, size=12, east_asia="宋体")
            set_paragraph_line_spacing(p, before=2, after=2, line=1.4)
            continue

        nm = re.match(r"^\s*\d+[.)]\s+(.*)$", line)
        if nm:
            list_serial += 1
            body = clean_inline_md(nm.group(1))
            p = doc.add_paragraph()
            r = p.add_run(f"{list_serial}. {body}")
            set_run_font(r, size=12, east_asia="宋体")
            set_paragraph_line_spacing(p, before=2, after=2, line=1.4)
            p.paragraph_format.left_indent = Pt(24)
            p.paragraph_format.first_line_indent = Pt(-24)
            continue

        p = doc.add_paragraph()
        r = p.add_run(clean_inline_md(line))
        set_run_font(r, size=12, east_asia="宋体")
        set_paragraph_line_spacing(p, before=2, after=2, line=1.6)
        p.paragraph_format.first_line_indent = Pt(24)


def main():
    parser = argparse.ArgumentParser(description="将 PRD Markdown 导出为印刷版 Word（封面 + 目录域 + 正文）")
    parser.add_argument(
        "--source",
        type=Path,
        default=DEFAULT_SOURCE_MD,
        help=f"源 Markdown 路径（默认：{DEFAULT_SOURCE_MD.name}）",
    )
    parser.add_argument(
        "--output",
        type=Path,
        default=DEFAULT_TARGET_DOCX,
        help=f"输出 docx 路径（默认：{DEFAULT_TARGET_DOCX.name}）",
    )
    parser.add_argument("--version", default=DEFAULT_COVER_VERSION, help="封面版本号，如 v2.0")
    parser.add_argument("--date", default=DEFAULT_COVER_DATE, dest="cover_date", help='封面编制日期，如 "2026年5月"')
    parser.add_argument(
        "--copy-final-alias",
        action="store_true",
        help=f"额外复制一份到项目约定的终稿长文件名（仅 v1.0 印刷版惯例使用）",
    )
    args = parser.parse_args()

    source_md = args.source if args.source.is_absolute() else _ROOT / args.source
    target_docx = args.output if args.output.is_absolute() else _ROOT / args.output

    if not source_md.exists():
        raise FileNotFoundError(f"找不到源文件: {source_md}")

    lines = source_md.read_text(encoding="utf-8").splitlines()

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(12)

    for hi in range(1, 10):
        try:
            hs = doc.styles[f"Heading {hi}"]
            hs.font.color.rgb = _HEADING_BLACK
        except KeyError:
            pass

    build_cover(doc, version=args.version, cover_date=args.cover_date)

    # 新节：目录与正文（从这里开始启用页眉页脚）
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    apply_header_footer(section)

    toc_title = doc.add_paragraph("目录")
    toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = toc_title.runs[0]
    set_run_font(tr, size=18, bold=True, east_asia="黑体")
    set_paragraph_line_spacing(toc_title, before=12, after=12, line=1.2)

    toc_p = doc.add_paragraph()
    add_toc(toc_p)
    for run in toc_p.runs:
        set_run_font(run, size=12, east_asia="宋体")

    doc.add_section(WD_SECTION.NEW_PAGE)
    apply_header_footer(doc.sections[-1])

    add_content(doc, lines)
    doc.save(target_docx)
    print(f"生成成功: {target_docx}")
    if args.copy_final_alias:
        shutil.copy2(target_docx, TARGET_DOCX_FINAL_ALIAS)
        print(f"已同步终稿文件名: {TARGET_DOCX_FINAL_ALIAS}")


if __name__ == "__main__":
    main()
